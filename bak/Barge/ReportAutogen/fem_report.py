from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# import openpyxl as xl

document = Document()

# wb = xl.load_workbook(filename = 'project_summary.xlsx')
# sht = wb['Project']

# project_name = sht.cell(row = 2, column = 2).value
# module_name = sht.cell(row = 3, column = 2).value
# rule_name = sht.cell(row = 4, column = 2).value

#document.add_paragraph(    'GENERAL', style='List Number')
#document.add_paragraph(    'Introduction', style='List Number 2')

document.add_paragraph('일반', style='Heading 1')

f = open('0일반.txt', 'r')
General = f.read()
document.add_paragraph('개요',style='Heading 2')
p = document.add_paragraph(General, style='Normal')
f.close()

f = open('0일반지침.txt', 'r')
GuideLine = f.read()
document.add_paragraph('일반지침', style='Heading 2')
p = document.add_paragraph(GuideLine, style='Normal')
f.close()

f = open('Limitation.txt', 'r')
Limitation = f.read()
document.add_paragraph('한계치',style='Heading 2')
p = document.add_paragraph(Limitation, style='Normal')
f.close()

f = open('Principal.txt', 'r')
principal = f.read()
document.add_paragraph('주요 사항',style='Heading 2')
p = document.add_paragraph(principal, style='Normal')
f.close()

f = open('0약어.txt', 'r')
abbreviations = f.read()
document.add_paragraph('약어',style='Heading 2')
p = document.add_paragraph(abbreviations, style='Normal')
f.close()

##################################################################################
document.add_paragraph('단위계',style='List Number 2')
unit_conversion = (
(0.03937, 'MILIMETERS', 'INCHES', 25.400 ),
(0.3937, 'CENTIMETERS', 'INCHES', 2.540 ),
(3.2808, 'METERS', 'FEET', 0.3048 ),
(2.2046, 'KILOGRAMMES', 'POUNDS', 0.45359 ),
(0.0009842, 'KILOGRAMMES', 'TONS(2240 lbs)', 1016.047),
(0.9842, 'METRIC TONS(ie TONNES OF 1000) KILOS', 'TONS(2240 lbs)', 1.016 ),
(2.4998, 'METRIC TONS PER CENTIMETER (OF IMMERSION)', 'TONS PER INCH(IMMERSION)', 0.400 ),
(8.2014, 'MOMENT TO CHANGE TRIM ONE CENTIMETER(TONESS METER UNIT)', 'MOMENT TO CHANGE(FOOT TON UNITS)', 0.122 ),
(187.9767, 'METER RADIANS', 'FEET DEGREE', 0.0053 )
)

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'MULTIPLY BY'
hdr_cells[1].text = 'TO CONVERT FROM'
hdr_cells[2].text = 'TO OBTAIN'
for co1, co2, co3, co4 in unit_conversion:
    row_cells = table.add_row().cells
    row_cells[0].text = str(co1)
    row_cells[1].text = co2
    row_cells[2].text = co3
    row_cells[3].text = str(co4)

#table.style = 'Light Shading'
table.style = 'Light Grid Accent 1'
document.add_page_break()
###############################################################################
f = open('0계산기본자료.txt', 'r')
basic_data = f.read()
document.add_paragraph('계산 기본 자료',style='List Number 2')
p = document.add_paragraph(basic_data)
f.close()

f = open('0변위및무게중심.txt', 'r')
cog = f.read()
document.add_paragraph('변위와 무게중심의 계산', style='List Number 2')
p = document.add_paragraph(cog)
f.close()

f = open('0트림과흘수계산.txt', 'r')
trim = f.read()
document.add_paragraph('하중에 대한 트림과 흘수의 계산', style='List Number 2')
p = document.add_paragraph(trim)
f.close()

document.add_picture('ship.png', width=Inches(1.25*5))

f = open('0메타센터계산.txt', 'r')
GOM = f.read()
document.add_paragraph('메타센터 높이계산(GoM', style='List Number 2')
p = document.add_paragraph(GOM)
f.close()
document.add_page_break()
#12
##################################################################################
document.add_paragraph('하중 중량에 따른 흘수와 트림',style='List Number 2')
document.add_paragraph('(Deflection은 고려되지 않는다.)')

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'WEIGHT'
hdr_cells[2].text = 'L.C.G'
hdr_cells[3].text = 'L-MOMENT'

row_cells = table.add_row().cells
row_cells[0].text = 'DISPLACEMENT'
row_cells[1].text = '14'
row_cells[2].text = '18'
row_cells[3].text = '14x18'

row_cells = table.add_row().cells
row_cells[0].text = 'LOADING WEIGHT'
row_cells[1].text = ''
row_cells[2].text = ''
row_cells[3].text = ''

row_cells = table.add_row().cells
row_cells[0].text = 'UNLOADING WEIGHT'
row_cells[1].text = ''
row_cells[2].text = ''
row_cells[3].text = ''

row_cells = table.add_row().cells
row_cells[0].text = 'TOTAL'
row_cells[1].text = '19'
row_cells[2].text = '25'
row_cells[3].text = ''
table.style = 'Light Grid Accent 1'      
##################################################################################
#########################
#########################

#13

##################################################################################
document.add_paragraph('COG와 트림 계산',style='List Number 2')

table = document.add_table(rows=15, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ITEM'
hdr_cells[1].text = 'WEIGHT(TON)'
hdr_cells[2].text = 'L.C.G(M)'
hdr_cells[3].text = 'L.C.G-MT(TON-M)'                       
hdr_cells[4].text = 'K.G'
hdr_cells[5].text = 'K.G-MT(TON-M)' 
                       
row_cells = table.rows[14].cells
row_cells[0].text = 'TOTAL'
table.style = 'Light Grid Accent 1'      
##################################################################################
    
document.add_paragraph('')
##################################################################################
table = document.add_table(rows=15, cols=1)
table.style = 'Light Grid Accent 1'      
##################################################################################

document.add_paragraph('')
##################################################################################
table = document.add_table(rows=15, cols=4)
table.style = 'Light Grid Accent 1'      
##################################################################################

f01 = open('../DBHULL.OUT', 'r')
f02 = open('../LOADCD.OUT', 'r')
f03 = open('../HYDROT.OUT', 'r')

f10 = f01.read()
f20 = f02.read()
f30 = f03.read()

f10 = f10.replace('','\n')
f20 = f20.replace('','\n')
f30 = f30.replace('','\n')
document.add_page_break()
p = document.add_paragraph(f10)
document.add_page_break()
p = document.add_paragraph(f20)
document.add_page_break()
p = document.add_paragraph(f30)

f01.close()
f02.close()
f03.close()
################################################
fldChar= OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText= OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text= r'TOC \f \o &amp;quot;1-9&amp;quot; \h' # change 1-3 depending on heading levels you need
fldChar2= OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3= OxmlElement('w:t')
fldChar3.text= "Right-click to update field."
fldChar2.append(fldChar3)
fldChar4= OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
r_element= run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)


########################################################
'''
# introduction = introduction.replace('BBBBB',module_name)
# introduction = introduction.replace('DDDDD',rule_name)



f101 = f10.read()
f111 = f11.read()
f121 = f12.read()



p = document.add_paragraph(f101)
p = document.add_paragraph(f111)
p = document.add_paragraph(f121)

f10.close()
f11.close()
f12.close()

p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
'''

document.save('demo.docx')
