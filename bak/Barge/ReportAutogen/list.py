from docx import Document
from docx.shared import Inches

document.add_paragraph(
    'GENERAL', style='List Number'
)


document.add_paragraph(
    'Introduction', style='List Number 2'
)

f = open('introduction.txt', 'r')

introduction = f.read()

introduction = introduction.replace('AAAAA',project_name)
introduction = introduction.replace('BBBBB',module_name)
introduction = introduction.replace('DDDDD',rule_name)

    
p = document.add_paragraph(introduction)

f.close()

p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.save('demo.docx')
